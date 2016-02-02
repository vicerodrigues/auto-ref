#!/usr/bin/env python3

import docx

doc=docx.Document('teste2.docx')

for parag in range(len(doc.paragraphs)):
    for myrun in range(len(doc.paragraphs[parag].runs)):
        if doc.paragraphs[parag].runs[myrun].text.startswith('[') and\
 doc.paragraphs[parag].runs[myrun].text.endswith(']'):
            doc.paragraphs[parag].runs[myrun].text='Adicionei um\
 texto aqui.'
doc.save('teste3.docx')
